"""DOCX document information extraction tool.

This module extracts structured information from DOCX documents using OpenAI or
Azure OpenAI structured outputs. It supports extracting multiple records from a
single document (e.g., multiple rows in a table) and batch processing multiple
documents in a folder.

Features:
    - Single file processing: Extract structured information from a single DOCX
    - Batch processing: Process all DOCX files in a folder
    - Multi-format export: Support JSON and CSV format outputs
    - Multi-record extraction: Automatically identify multiple rows of data
    - Dual API support: Support both Azure OpenAI and standard OpenAI API

Typical usage examples:
    # Use Azure OpenAI (default)
    python main.py document.docx
    python main.py document.docx -o output.json

    # Batch process folder and export to CSV (also creates results.json)
    python main.py ./documents -o results.csv

    # Convert JSON results back to DOCX with nested tables
    python main.py --json-to-docx results.json -o output.docx

    # Use custom Azure configuration
    python main.py document.docx --api-key your-azure-key --model your-deployment-name

    # Use standard OpenAI API
    set USE_AZURE_OPENAI=false
    python main.py document.docx --api-key your-openai-key

Environment variables (Azure OpenAI - enabled by default):
    USE_AZURE_OPENAI: Whether to use Azure OpenAI (default: true)
    AZURE_OPENAI_API_KEY: Azure OpenAI API key
    AZURE_OPENAI_ENDPOINT: Azure OpenAI endpoint URL
    AZURE_OPENAI_API_VERSION: Azure OpenAI API version (default: 2024-08-01-preview)
    AZURE_OPENAI_DEPLOYMENT: Azure OpenAI deployment name (default: gpt-4o)

Environment variables (Standard OpenAI - when USE_AZURE_OPENAI=false):
    OPENAI_API_KEY: OpenAI API key
    OPENAI_API_BASE: OpenAI API base URL
    OPENAI_MODEL: Model name to use (default: gpt-4o-2024-08-06)

Note:
    - Each row of data in the document will be extracted as a separate record
    - When batch processing folders, CSV format is recommended for consolidation
    - CSV output includes source filename for tracking record origins
    - When exporting to CSV, a JSON file with the same name is also automatically created
    - Nested tables are preserved using special markers (<<NESTED_TABLE_N_START>> ... <<NESTED_TABLE_N_END>>)
    - Use --json-to-docx to convert JSON results back to DOCX with actual nested tables
    - Azure OpenAI is used by default, set USE_AZURE_OPENAI=false to switch to standard OpenAI
"""

import argparse
import asyncio
import csv
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import AzureOpenAI, OpenAI, AsyncAzureOpenAI, AsyncOpenAI
from pydantic import BaseModel, Field
from tqdm import tqdm


class DocumentFields(BaseModel):
    """Structured model for document fields.

    Attributes:
        tl_ea: TL EA information from Column 1 of attached protocol.
        test_standard: Test standard from Column 2 (excluding website links).
        test_analytes: Test analytes from Column 5.
        pp_notes: PP notes information from Column 3.
        source_link: Website link from Column 2 if found, otherwise None.
        label_and_symbol: Whether any label is found in this row (yes/no).
    """
    tl_ea: str = Field(description="Column 1 of attached protocol - TL EA information")
    test_standard: str = Field(description="Column 2 but not website - test standard (excluding website links)")
    test_analytes: str = Field(description="Column 5 - test analytes")
    pp_notes: str = Field(description="Column 3 - PP notes information")
    source_link: Optional[str] = Field(default=None, description="Column 2 website if found - source link")
    label_and_symbol: str = Field(description="Any label found in this row, just state yes/no")


class DocumentExtraction(BaseModel):
    """Document extraction result model containing multiple records.

    Attributes:
        records: List of all records extracted from the document.
    """
    records: list[DocumentFields] = Field(
        description="All records extracted from the document, each record corresponds to a row of data"
    )


class DocxExtractor:
    """DOCX document extractor for extracting structured information from Word documents.

    This class uses OpenAI or Azure OpenAI structured output functionality to
    extract specific fields from DOCX documents.

    Attributes:
        client: OpenAI or AzureOpenAI client instance.
        model: Name of the OpenAI model or Azure deployment to use.
        use_azure: Whether Azure OpenAI is being used.
    """

    def __init__(
        self,
        api_key: str,
        model: str = "gpt-5-chat",
        api_base: Optional[str] = None,
        use_azure: bool = False,
        azure_endpoint: Optional[str] = None,
        azure_api_version: Optional[str] = None
    ):
        """Initializes a DocxExtractor instance.

        Args:
            api_key: OpenAI or Azure OpenAI API key.
            model: Model name or Azure deployment name. Defaults to "gpt-4o-2024-08-06".
            api_base: API base URL. Defaults to None (uses OpenAI official address).
            use_azure: Whether to use Azure OpenAI. Defaults to False.
            azure_endpoint: Azure OpenAI endpoint URL (required when use_azure=True).
            azure_api_version: Azure OpenAI API version (required when use_azure=True).

        Raises:
            openai.OpenAIError: If the API key is invalid.
            ValueError: If using Azure but missing required parameters.
        """
        if use_azure:
            if not azure_endpoint:
                raise ValueError("azure_endpoint is required when using Azure OpenAI")
            if not azure_api_version:
                raise ValueError("azure_api_version is required when using Azure OpenAI")

            self.client = AzureOpenAI(
                api_key=api_key,
                azure_endpoint=azure_endpoint,
                api_version=azure_api_version
            )
            # Create async client for batch processing
            self.async_client = AsyncAzureOpenAI(
                api_key=api_key,
                azure_endpoint=azure_endpoint,
                api_version=azure_api_version
            )
        else:
            client_kwargs = {"api_key": api_key}
            if api_base:
                client_kwargs["base_url"] = api_base

            self.client = OpenAI(**client_kwargs)
            # Create async client for batch processing
            self.async_client = AsyncOpenAI(**client_kwargs)

        self.model = model
        self.use_azure = use_azure

    def read_docx(self, file_path: str) -> str:
        """Reads a DOCX file and extracts all text content.

        This method extracts paragraph text and table content from the document,
        combining them into a single string. It also detects and marks cells
        that contain images or nested tables. Only processes 5-column tables.

        Args:
            file_path: Path to the DOCX file (relative or absolute).

        Returns:
            A string containing all document content, with paragraphs and tables
            separately annotated.

        Raises:
            FileNotFoundError: If the specified file does not exist.
            docx.opc.exceptions.PackageNotFoundError: If the file is not a valid DOCX format.
        """
        doc = Document(file_path)

        # Extract all paragraph text
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extract table content with image and nested table detection
        tables_content = []
        for table in doc.tables:
            # Skip tables that are not 5-column tables
            if not self._is_valid_main_table(table):
                print(f"  Skipping table with {len(table.rows[0].cells) if table.rows else 0} columns (expected 5)")
                continue

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Check if cell contains nested tables
                    if self._cell_has_nested_table(cell):
                        # Extract content with tables in their original positions
                        cell_content = self._extract_cell_content_with_tables(cell)
                    else:
                        # No nested tables, just get the text
                        cell_content = cell.text.strip()

                    # Check if cell contains images
                    if self._cell_has_images(cell):
                        if cell_content:
                            cell_content += " [IMAGE: Contains label/symbol image]"
                        else:
                            cell_content = "[IMAGE: Contains label/symbol image]"

                    row_data.append(cell_content)
                if any(row_data):  # Only add non-empty rows
                    tables_content.append(" | ".join(row_data))

        # Combine all content
        full_text = "\n".join(paragraphs)
        if tables_content:
            full_text += "\n\n=== Table Content ===\n" + "\n".join(tables_content)

        return full_text

    def _cell_has_images(self, cell) -> bool:
        """Check if a table cell contains any images.

        Args:
            cell: A docx table cell object.

        Returns:
            True if the cell contains one or more images, False otherwise.
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                # Check if the run contains any drawing elements (images)
                if run._element.xpath('.//w:drawing'):
                    return True
                # Also check for inline shapes (another way images can be embedded)
                if run._element.xpath('.//w:pict'):
                    return True
        return False

    def _cell_has_nested_table(self, cell) -> bool:
        """Check if a table cell contains nested tables.

        Args:
            cell: A docx table cell object.

        Returns:
            True if the cell contains one or more nested tables, False otherwise.
        """
        # Use cell.tables property to check for nested tables
        return len(cell.tables) > 0

    def _format_nested_table_as_markdown(self, table, table_index: int = 0) -> str:
        """Format a single nested table as a Markdown table with special delimiters.

        Args:
            table: A docx table object.
            table_index: Index of the nested table (for unique identification).

        Returns:
            Markdown-formatted table string with special delimiters.
        """
        markdown_rows = []

        # Process all rows
        for row_idx, row in enumerate(table.rows):
            row_data = [c.text.strip() for c in row.cells]

            if any(row_data):  # Only add non-empty rows
                # Format as Markdown table row
                markdown_rows.append('| ' + ' | '.join(row_data) + ' |')

                # Add separator after first row (header)
                if row_idx == 0 and len(row_data) > 0:
                    markdown_rows.append('| ' + ' | '.join(['---'] * len(row_data)) + ' |')

        if markdown_rows:
            # Wrap with special delimiters
            return f'<<NESTED_TABLE_{table_index}_START>>\n' + '\n'.join(markdown_rows) + f'\n<<NESTED_TABLE_{table_index}_END>>'
        return ''

    def _extract_cell_content_with_tables(self, cell) -> str:
        """Extract cell content with nested tables in their original positions.

        This method preserves the order of text and tables as they appear in the cell.
        Nested tables are converted to Markdown format with special delimiters and
        inserted at their original positions in the text flow.

        Args:
            cell: A docx table cell object.

        Returns:
            A string containing the cell content with nested tables marked and
            positioned correctly in the original order.
        """
        content_parts = []
        table_index = 0

        # Iterate through all elements in the cell in order
        for element in cell._element:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                # Get paragraph text
                para_text = ''.join([
                    node.text for node in element.xpath('.//w:t') if node.text
                ]).strip()
                if para_text:
                    content_parts.append(para_text)

            # Check if it's a table
            elif element.tag.endswith('tbl'):
                # Find the corresponding table object
                for table in cell.tables:
                    if table._element == element:
                        markdown_table = self._format_nested_table_as_markdown(table, table_index)
                        if markdown_table:
                            content_parts.append(markdown_table)
                        table_index += 1
                        break

        # Join all parts with appropriate spacing
        return '\n\n'.join(content_parts) if content_parts else cell.text.strip()

    def _is_valid_main_table(self, table) -> bool:
        """Check if a table is a valid 5-column main table to process.

        Args:
            table: A docx table object.

        Returns:
            True if the table has 5 columns, False otherwise.
        """
        if len(table.rows) == 0:
            return False

        # Check the number of columns from the first row (header row)
        first_row = table.rows[0]
        num_columns = len(first_row.cells)

        return num_columns == 5

    @staticmethod
    def estimate_tokens(text: str) -> int:
        """Estimates the number of tokens in a text string.

        This uses a simple heuristic:
        - For ASCII characters: ~4 characters per token
        - For non-ASCII (including Chinese): ~2 characters per token

        Args:
            text: The text to estimate tokens for.

        Returns:
            Estimated number of tokens.
        """
        if not text:
            return 0

        ascii_chars = sum(1 for c in text if ord(c) < 128)
        non_ascii_chars = len(text) - ascii_chars

        # Rough estimation: 4 chars/token for ASCII, 2 chars/token for non-ASCII
        estimated_tokens = (ascii_chars / 4.0) + (non_ascii_chars / 2.0)

        return int(estimated_tokens) + 1  # Add 1 for safety margin

    def read_docx_in_batches(
        self,
        file_path: str,
        batch_size: int = 10,
        max_tokens_per_batch: int = 20000
    ) -> list[str]:
        """Reads a DOCX file and splits table content into batches by rows with dynamic token window.

        This method extracts table content and splits it into multiple batches.
        It uses a dynamic window strategy that considers both row count and token limits
        to ensure the output doesn't get truncated due to exceeding token limits.

        Args:
            file_path: Path to the DOCX file (relative or absolute).
            batch_size: Maximum number of data rows per batch. Defaults to 10.
                This serves as an upper limit; actual batch size may be smaller
                if token limit is reached.
            max_tokens_per_batch: Maximum tokens per batch input. Defaults to 20000.
                This ensures enough room for model output (typically ~4-8k tokens).

        Returns:
            A list of strings, each containing a batch of table content with header.

        Raises:
            FileNotFoundError: If the specified file does not exist.
            docx.opc.exceptions.PackageNotFoundError: If the file is not a valid DOCX format.
        """
        doc = Document(file_path)
        batches = []

        # Extract paragraph text (will be prepended to each batch)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        paragraph_text = "\n".join(paragraphs) if paragraphs else ""

        # Reserve tokens for system prompt (~2000), paragraph text, and formatting
        system_prompt_tokens = 2000
        paragraph_tokens = self.estimate_tokens(paragraph_text)
        reserved_tokens = system_prompt_tokens + paragraph_tokens + 500  # 500 for formatting

        # Process each table
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) == 0:
                continue

            # Skip tables that are not 5-column tables
            if not self._is_valid_main_table(table):
                print(f"  Skipping table {table_idx + 1} with {len(table.rows[0].cells)} columns (expected 5)")
                continue

            # Extract header row with image and nested table detection
            header_row = table.rows[0]
            header_data = []
            for cell in header_row.cells:
                # Check if cell contains nested tables
                if self._cell_has_nested_table(cell):
                    # Extract content with tables in their original positions
                    cell_content = self._extract_cell_content_with_tables(cell)
                else:
                    # No nested tables, just get the text
                    cell_content = cell.text.strip()

                # Check if cell contains images
                if self._cell_has_images(cell):
                    if cell_content:
                        cell_content += " [IMAGE: Contains label/symbol image]"
                    else:
                        cell_content = "[IMAGE: Contains label/symbol image]"

                header_data.append(cell_content)

            header_line = " | ".join(header_data)
            header_tokens = self.estimate_tokens(header_line)

            # Extract data rows with image and nested table detection
            data_rows = []
            data_row_tokens = []
            for row in table.rows[1:]:  # Skip header
                row_data = []
                for cell in row.cells:
                    # Check if cell contains nested tables
                    if self._cell_has_nested_table(cell):
                        # Extract content with tables in their original positions
                        cell_content = self._extract_cell_content_with_tables(cell)
                    else:
                        # No nested tables, just get the text
                        cell_content = cell.text.strip()

                    # Check if cell contains images
                    if self._cell_has_images(cell):
                        if cell_content:
                            cell_content += " [IMAGE: Contains label/symbol image]"
                        else:
                            cell_content = "[IMAGE: Contains label/symbol image]"

                    row_data.append(cell_content)
                if any(row_data):  # Only add non-empty rows
                    row_line = " | ".join(row_data)
                    data_rows.append(row_line)
                    data_row_tokens.append(self.estimate_tokens(row_line))

            # Dynamic batching based on token count
            batch_idx = 0
            i = 0
            while i < len(data_rows):
                batch_data_rows = []
                current_batch_tokens = reserved_tokens + header_tokens

                # Add rows until we hit batch_size or token limit
                rows_in_batch = 0
                while i < len(data_rows) and rows_in_batch < batch_size:
                    row_tokens = data_row_tokens[i]

                    # Check if adding this row would exceed token limit
                    if current_batch_tokens + row_tokens > max_tokens_per_batch:
                        # If this is the first row in batch, we must include it anyway
                        if rows_in_batch == 0:
                            batch_data_rows.append(data_rows[i])
                            i += 1
                            print(f"Warning: Single row exceeds token limit ({row_tokens} tokens)")
                        break

                    batch_data_rows.append(data_rows[i])
                    current_batch_tokens += row_tokens
                    rows_in_batch += 1
                    i += 1

                # Construct batch content
                if batch_data_rows:
                    batch_content = []
                    if paragraph_text:
                        batch_content.append(paragraph_text)

                    batch_content.append(f"\n=== Table {table_idx + 1} (Batch {batch_idx + 1}) ===")
                    batch_content.append(header_line)
                    batch_content.extend(batch_data_rows)

                    batch_text = "\n".join(batch_content)
                    batches.append(batch_text)

                    # Log batch info
                    actual_tokens = self.estimate_tokens(batch_text)
                    print(f"  Batch {batch_idx + 1}: {rows_in_batch} rows, ~{actual_tokens} tokens")

                    batch_idx += 1

        # If no tables found, return the whole document as one batch
        if not batches:
            batches.append(self.read_docx(file_path))

        return batches

    def convert_tables_to_markdown(self, file_path: str) -> str:
        """Converts all 5-column tables in a DOCX file to Markdown format.

        This method extracts all 5-column tables from the document and converts them
        to Markdown table syntax. Non-table content is preserved as plain text.
        It also detects and marks cells that contain images or nested tables.

        Args:
            file_path: Path to the DOCX file (relative or absolute).

        Returns:
            A string containing the document content with tables in Markdown format.

        Raises:
            FileNotFoundError: If the specified file does not exist.
            docx.opc.exceptions.PackageNotFoundError: If the file is not a valid DOCX format.
        """
        doc = Document(file_path)
        output = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)
                output.append("")  # Add blank line

        # Extract and convert tables to Markdown
        valid_table_count = 0
        for table_idx, table in enumerate(doc.tables, 1):
            # Skip tables that are not 5-column tables
            if not self._is_valid_main_table(table):
                output.append(f"### Table {table_idx} (Skipped - {len(table.rows[0].cells) if table.rows else 0} columns, expected 5)")
                output.append("")
                continue

            valid_table_count += 1
            output.append(f"### Table {table_idx} (5 columns)")
            output.append("")

            # Get table data with image and nested table detection
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Check if cell contains nested tables
                    if self._cell_has_nested_table(cell):
                        # Extract content with tables in their original positions
                        cell_content = self._extract_cell_content_with_tables(cell)
                    else:
                        # No nested tables, just get the text
                        cell_content = cell.text.strip()

                    # Check if cell contains images
                    if self._cell_has_images(cell):
                        if cell_content:
                            cell_content += " [IMAGE: Contains label/symbol image]"
                        else:
                            cell_content = "[IMAGE: Contains label/symbol image]"

                    row_data.append(cell_content)
                table_data.append(row_data)

            if not table_data:
                continue

            # Assume first row is header
            if len(table_data) > 0:
                header = table_data[0]
                num_cols = len(header)

                # Write header
                output.append("| " + " | ".join(header) + " |")

                # Write separator
                output.append("| " + " | ".join(["---"] * num_cols) + " |")

                # Write data rows
                for row in table_data[1:]:
                    # Pad row if it has fewer columns than header
                    while len(row) < num_cols:
                        row.append("")
                    output.append("| " + " | ".join(row[:num_cols]) + " |")

                output.append("")  # Add blank line after table

        return "\n".join(output)

    async def extract_fields_async(self, text: str) -> DocumentExtraction:
        """Asynchronously extracts document fields using OpenAI structured output API.

        This method is used for all extraction operations (batch processing).

        Args:
            text: Text content extracted from the DOCX document.

        Returns:
            A DocumentExtraction instance containing all extracted records.

        Raises:
            openai.APIError: If the API call fails.
            openai.RateLimitError: If the API rate limit is exceeded.
        """
        completion = await self.async_client.responses.parse(
            model=self.model,
            input=[
                {
                    "role": "system",
                    "content": """You are a professional document information extraction assistant. Please extract the following fields from the provided document content:

                    IMPORTANT - Table Identification:
                    - The document may contain multiple tables, but you should ONLY process the main 5-column data table
                    - The main target table has 5 columns: Column 1 (TL EA), Column 2 (Test Standard), Column 3 (PP Notes), Column 4, and Column 5 (Test Analytes)
                    - IGNORE any other tables that are not the main 5-column table (e.g., auxiliary tables, summary tables, etc.)
                    - If you encounter content from non-target tables in the input, skip them completely - do not create records for them
                    - Only create records for rows that clearly belong to the main 5-column data table
                    - When in doubt, check if the row has the expected 5-column structure with the fields described below

                    Field Extraction Instructions:
                    1. TL EA: Extract the attached protocol information from Column 1
                       - Extract ALL content completely and accurately from Column 1
                       - Do not omit, summarize, or truncate any text
                       - If Column 1 contains a nested table (marked with "<<NESTED_TABLE_N_START>>" ... "<<NESTED_TABLE_N_END>>"), PRESERVE the complete marker format including the delimiters and all content between them exactly as they appear
                    2. Test standard: Extract non-website content from Column 2 (test standard)
                       - Extract ALL text content from Column 2 completely
                       - Exclude only website URLs (which should go to Source link field)
                       - Do not omit, summarize, or truncate any text
                       - If Column 2 contains a nested table (marked with "<<NESTED_TABLE_N_START>>" ... "<<NESTED_TABLE_N_END>>"), PRESERVE the complete marker format including the delimiters and all content between them exactly as they appear
                    3. Test analytes: Extract test analyte information ONLY from Column 5
                       - IMPORTANT: Only extract content that appears in Column 5
                       - Do NOT extract test analytes or chemical names from Column 3 (PP notes) or any other columns
                       - If Column 5 is empty, return an empty string
                       - Only use the exact content from Column 5, do not infer or analyze from other columns
                       - Extract ALL content from Column 5 completely when present
                       - If Column 5 contains a nested table (marked with "<<NESTED_TABLE_N_START>>" ... "<<NESTED_TABLE_N_END>>"), PRESERVE the complete marker format including the delimiters and all content between them exactly as they appear
                    4. PP notes: Extract notes information from Column 3
                       - CRITICAL: Extract ALL content from Column 3 completely and accurately
                       - Do NOT omit any text even if it contains chemical names or test information
                       - Do NOT summarize or truncate the content
                       - Column 3 often contains detailed requirements, standards, and notes - preserve everything
                       - If Column 3 contains a nested table (marked with "<<NESTED_TABLE_N_START>>" ... "<<NESTED_TABLE_N_END>>"), PRESERVE the complete marker format including the delimiters and all content between them exactly as they appear
                    5. Source link: If there is a website link in Column 2, extract it; otherwise return null
                    6. Label and symbol: This field indicates whether this row contains any label or symbol images (such as certification marks, safety labels, warning symbols, etc.).
                       - In the document, cells containing images will be marked with "[IMAGE: Contains label/symbol image]"
                       - If you see this marker in any cell of the current row, return "yes"
                       - If no such marker is found in the row, return "no"
                       - Note: The label/symbol refers to visual graphics or icons embedded in the document, not just text descriptions

                    Important notes about nested tables:
                    - Some cells may contain nested tables, marked with special delimiters: "<<NESTED_TABLE_N_START>>" and "<<NESTED_TABLE_N_END>>" where N is the table index
                    - The nested table content between delimiters is formatted as Markdown tables (with | separators and --- header dividers)
                    - Multiple nested tables in the same cell will have different indices (0, 1, 2, etc.)
                    - CRITICAL: You MUST preserve these special markers EXACTLY as they appear in your output
                    - DO NOT parse, interpret, summarize, or convert the nested tables - just copy them verbatim with their delimiters
                    - The markers are special tokens that will be processed later to recreate actual nested tables in the final output
                    - Include the complete content: start delimiter, Markdown table content, and end delimiter
                    - Example: If you see "<<NESTED_TABLE_0_START>>\n| Header1 | Header2 |\n| --- | --- |\n| Data1 | Data2 |\n<<NESTED_TABLE_0_END>>", you must output it EXACTLY as is in your extraction result
                    - The main document only processes 5-column tables - any other tables are nested within cells

                    Important notes:
                    - The document may contain multiple rows of data (e.g., multiple rows in a table)
                    - Please create a separate record for each row of data
                    - Put all records in the records list
                    - Please carefully analyze the document content and accurately extract this information
                    - Pay special attention to the "[IMAGE: Contains label/symbol image]" markers to determine the Label and symbol field
                    - CRITICAL: PRESERVE all "<<NESTED_TABLE_N_START>>" ... "<<NESTED_TABLE_N_END>>" markers EXACTLY as they appear
                    - DO NOT parse, interpret, or convert the nested table markers - copy them verbatim
                    - The markers (<<NESTED_TABLE_N_START>> and <<NESTED_TABLE_N_END>>) are special tokens that must be preserved exactly in your output
                    - CRITICAL: Test analytes must ONLY come from Column 5, never from other columns
                    - CRITICAL: All other columns (1, 2, 3) must be extracted COMPLETELY without omission, summarization, or truncation"""
                },
                {
                    "role": "user",
                    "content": f"Please extract information from all rows in the following document content:\n\n{text}"
                }
            ],
            text_format=DocumentExtraction,
        )

        return completion.output_parsed

    async def process_batches_async(
        self,
        batches: list[str],
        max_concurrent: int = 5
    ) -> DocumentExtraction:
        """Processes multiple batches of document content asynchronously with concurrency control.

        This method processes multiple batches concurrently using asyncio, with a limit
        on the number of concurrent requests to avoid overwhelming the API.

        Args:
            batches: List of text batches to process.
            max_concurrent: Maximum number of concurrent API calls. Defaults to 5.

        Returns:
            A DocumentExtraction instance containing all records from all batches,
            maintaining the order of batches.

        Raises:
            openai.APIError: If any API call fails.
        """
        semaphore = asyncio.Semaphore(max_concurrent)

        async def process_with_semaphore(batch_text: str, batch_idx: int) -> tuple[int, DocumentExtraction]:
            async with semaphore:
                result = await self.extract_fields_async(batch_text)
                return batch_idx, result

        # Create tasks for all batches with progress tracking
        tasks = [process_with_semaphore(batch, idx) for idx, batch in enumerate(batches)]

        # Execute all tasks with progress bar
        all_records = []
        with tqdm(total=len(batches), desc="Processing batches", unit="batch") as pbar:
            for coro in asyncio.as_completed(tasks):
                batch_idx, extraction = await coro
                all_records.append((batch_idx, extraction))
                pbar.update(1)

        # Sort by batch index to maintain order
        all_records.sort(key=lambda x: x[0])

        # Combine all records from all batches
        combined_records = []
        for _, extraction in all_records:
            combined_records.extend(extraction.records)

        return DocumentExtraction(records=combined_records)

    @staticmethod
    def _parse_nested_tables(text: str) -> list[tuple[str, str, str]]:
        """Parse nested table markers from text and return list of (start_marker, end_marker, content).

        Args:
            text: Text containing nested table markers.

        Returns:
            List of tuples containing (start_marker, end_marker, table_content).
        """
        pattern = r'<<NESTED_TABLE_(\d+)_START>>\n(.*?)\n<<NESTED_TABLE_\1_END>>'
        matches = re.finditer(pattern, text, re.DOTALL)

        result = []
        for match in matches:
            table_index = match.group(1)
            table_content = match.group(2)
            start_marker = f'<<NESTED_TABLE_{table_index}_START>>'
            end_marker = f'<<NESTED_TABLE_{table_index}_END>>'
            result.append((start_marker, end_marker, table_content))

        return result

    @staticmethod
    def _markdown_to_docx_table(doc_or_cell, markdown_table: str):
        """Convert a Markdown table string to a DOCX table.

        Args:
            doc_or_cell: A Document or table cell object where the table should be added.
            markdown_table: Markdown-formatted table string.

        Returns:
            The created table object.
        """
        # Parse markdown table
        lines = [line.strip() for line in markdown_table.strip().split('\n') if line.strip()]

        # Remove separator line (the one with ---)
        data_lines = [line for line in lines if not re.match(r'^\|\s*-+\s*(\|\s*-+\s*)*\|$', line)]

        # Parse rows
        rows_data = []
        for line in data_lines:
            # Remove leading/trailing pipes and split by pipe
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            rows_data.append(cells)

        if not rows_data:
            return None

        # Create table
        num_rows = len(rows_data)
        num_cols = len(rows_data[0]) if rows_data else 0

        table = doc_or_cell.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        # Fill table data
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_data
                    # Make header row bold
                    if i == 0:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

        return table

    @staticmethod
    def _add_text_with_nested_tables(cell, text: str):
        """Add text and nested tables to a cell, parsing special markers.

        Args:
            cell: A table cell object.
            text: Text that may contain nested table markers.
        """
        # Find all nested tables
        nested_tables = DocxExtractor._parse_nested_tables(text)

        if not nested_tables:
            # No nested tables, just add the text
            cell.text = text
            return

        # Split text by nested table markers and reconstruct
        remaining_text = text

        # Clear existing content
        cell.text = ''

        for start_marker, end_marker, table_content in nested_tables:
            # Find the position of this nested table
            full_marker = f'{start_marker}\n{table_content}\n{end_marker}'

            if full_marker in remaining_text:
                # Split by this marker
                before, _, after = remaining_text.partition(full_marker)

                # Add text before the table
                if before.strip():
                    cell.add_paragraph(before.strip())

                # Add the nested table
                DocxExtractor._markdown_to_docx_table(cell, table_content)

                # Continue with remaining text
                remaining_text = after

        # Add any remaining text after the last table
        if remaining_text.strip():
            cell.add_paragraph(remaining_text.strip())

    @staticmethod
    def export_to_docx(extraction: DocumentExtraction, output_path: str, source_file: Optional[str] = None) -> None:
        """Export extraction results to a DOCX file with nested tables.

        Args:
            extraction: DocumentExtraction object containing the extracted records.
            output_path: Output DOCX file path.
            source_file: Optional source filename to include in the header.

        Raises:
            IOError: If unable to write to the file.
        """
        doc = Document()

        # Add title
        title = doc.add_heading('Document Extraction Results', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add source file info if provided
        if source_file:
            info_para = doc.add_paragraph()
            info_para.add_run('Source File: ').bold = True
            info_para.add_run(source_file)
            doc.add_paragraph()  # Blank line

        # Create main table with headers
        main_table = doc.add_table(rows=1, cols=6)
        main_table.style = 'Table Grid'

        # Set header row
        header_cells = main_table.rows[0].cells
        headers = ['TL EA', 'Test Standard', 'Test Analytes', 'PP Notes', 'Source Link', 'Label and Symbol']

        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

        # Add data rows
        for record in extraction.records:
            row_cells = main_table.add_row().cells

            # Process each field and add to cell
            fields = [
                record.tl_ea,
                record.test_standard,
                record.test_analytes,
                record.pp_notes,
                record.source_link or '',
                record.label_and_symbol
            ]

            for i, field_value in enumerate(fields):
                if field_value:
                    DocxExtractor._add_text_with_nested_tables(row_cells[i], field_value)

        # Save document
        doc.save(output_path)

    @staticmethod
    def export_to_csv(extractions: list[tuple[str, DocumentExtraction]], output_path: str) -> None:
        """Exports multiple extraction results to a CSV file.

        Args:
            extractions: List of (filename, DocumentExtraction) tuples.
            output_path: Output CSV file path.

        Raises:
            IOError: If unable to write to the file.
        """
        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)

            # Write header
            writer.writerow([
                'Source File',
                'TL EA',
                'Test Standard',
                'Test Analytes',
                'PP Notes',
                'Source Link',
                'Label and Symbol'
            ])

            # Write all records from each file
            for filename, extraction in extractions:
                for record in extraction.records:
                    writer.writerow([
                        filename,
                        record.tl_ea,
                        record.test_standard,
                        record.test_analytes,
                        record.pp_notes,
                        record.source_link or '',
                        record.label_and_symbol
                    ])

    def process_file_with_batches(
        self,
        file_path: str,
        batch_size: int = 10,
        max_concurrent: int = 5,
        max_tokens_per_batch: int = 20000,
        output_path: Optional[str] = None
    ) -> DocumentExtraction:
        """Processes a DOCX file using batch processing with async API calls.

        This method splits the document into batches by table rows, then processes
        all batches concurrently using async API calls. Progress is shown via a
        progress bar. Uses dynamic token windowing to prevent output truncation.

        Args:
            file_path: Input DOCX file path (relative or absolute).
            batch_size: Maximum number of table rows per batch. Defaults to 10.
                Actual batch size may be smaller if token limit is reached.
            max_concurrent: Maximum number of concurrent API calls. Defaults to 5.
            max_tokens_per_batch: Maximum input tokens per batch. Defaults to 20000.
                This ensures enough room for model output without truncation.
            output_path: Optional output file path (.json or .csv). If provided, results
                will be saved in the specified format based on file extension.

        Returns:
            A DocumentExtraction instance containing all extracted records from all batches.

        Raises:
            FileNotFoundError: If the input file does not exist.
            openai.APIError: If the OpenAI API call fails.
            IOError: If unable to write to the output file.
        """
        print(f"Reading file: {file_path}")
        batches = self.read_docx_in_batches(file_path, batch_size, max_tokens_per_batch)

        print(f"Document split into {len(batches)} batch(es)")
        print(f"Processing with up to {max_concurrent} concurrent API calls...\n")

        # Run async processing
        extraction = asyncio.run(self.process_batches_async(batches, max_concurrent))

        print("\n\nExtraction complete!")
        print("=" * 80)
        print(f"Total records extracted: {len(extraction.records)}\n")

        for idx, record in enumerate(extraction.records, 1):
            print(f"Record #{idx}")
            print("-" * 80)
            print(f"  TL EA:           {record.tl_ea}")
            print(f"  Test standard:   {record.test_standard}")
            print(f"  Test analytes:   {record.test_analytes}")
            print(f"  PP notes:        {record.pp_notes}")
            print(f"  Source link:     {record.source_link}")
            print(f"  Label & symbol:  {record.label_and_symbol}")
            print()

        print("=" * 80)

        # Save based on file extension
        if output_path:
            output_path_obj = Path(output_path)
            filename = Path(file_path).name

            if output_path_obj.suffix.lower() == '.csv':
                # Export to CSV
                DocxExtractor.export_to_csv([(filename, extraction)], output_path)
                print(f"\nResults saved to CSV: {output_path}")

                # Also save JSON with same base name
                json_path = output_path_obj.with_suffix('.json')
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(extraction.model_dump(), f, ensure_ascii=False, indent=2)
                print(f"Results also saved to JSON: {json_path}")

            elif output_path_obj.suffix.lower() == '.docx':
                # Export to DOCX
                DocxExtractor.export_to_docx(extraction, output_path, filename)
                print(f"\nResults saved to DOCX: {output_path}")
            else:
                # Default to JSON
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(extraction.model_dump(), f, ensure_ascii=False, indent=2)
                print(f"\nResults saved to JSON: {output_path}")

        return extraction


def json_to_docx(json_path: str, output_path: str) -> None:
    """Convert a JSON extraction result file to DOCX format.

    Args:
        json_path: Path to the input JSON file.
        output_path: Path to the output DOCX file.

    Raises:
        FileNotFoundError: If the JSON file does not exist.
        json.JSONDecodeError: If the JSON file is malformed.
        IOError: If unable to write to the output file.
    """
    # Read JSON file
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Check if it's a single-file or multi-file JSON
    if 'files' in data:
        # Multi-file format - create a document for each file
        for file_data in data['files']:
            filename = file_data['filename']
            records = file_data['records']

            # Create extraction object
            extraction = DocumentExtraction(records=[DocumentFields(**r) for r in records])

            # Generate output filename
            base_name = Path(filename).stem
            output_file = Path(output_path).parent / f"{Path(output_path).stem}_{base_name}.docx"

            # Export to DOCX
            DocxExtractor.export_to_docx(extraction, str(output_file), filename)
            print(f"Exported {filename} to {output_file}")

    elif 'records' in data:
        # Single-file format
        extraction = DocumentExtraction(records=[DocumentFields(**r) for r in data['records']])

        # Get source filename from JSON path
        source_file = Path(json_path).stem

        # Export to DOCX
        DocxExtractor.export_to_docx(extraction, output_path, source_file)
        print(f"Exported to {output_path}")

    else:
        raise ValueError("Invalid JSON format: expected 'records' or 'files' key")


def parse_args() -> argparse.Namespace:
    """Parses command line arguments.

    Returns:
        A Namespace object containing the parsed arguments.
    """
    parser = argparse.ArgumentParser(
        description="DOCX Document Information Extraction Tool - Extract structured data from Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
OPERATION MODES:
  The tool has three main operation modes:

  1. EXTRACT MODE (default) - Extract information from DOCX files
     %(prog)s INPUT_FILE -o OUTPUT_FILE
     %(prog)s INPUT_FOLDER -o OUTPUT_FILE

  2. CONVERT MODE - Convert between formats without AI processing
     %(prog)s INPUT_FILE --mode convert --format markdown -o OUTPUT_FILE
     %(prog)s INPUT_JSON --mode convert --format docx -o OUTPUT_FILE

  3. Both modes support single file or batch folder processing

EXAMPLES:

  Basic Extraction:
    %(prog)s document.docx -o results.csv          # Extract to CSV+JSON
    %(prog)s document.docx -o results.json         # Extract to JSON only
    %(prog)s ./docs -o results.csv                 # Batch process folder

  Advanced Extraction (customize batch processing):
    %(prog)s document.docx -o results.csv --batch-size 20      # Adjust max rows per batch
    %(prog)s document.docx -o results.csv --concurrency 10     # Adjust concurrent API calls
    %(prog)s document.docx -o results.csv --max-tokens 25000   # Adjust token limit per batch

  Format Conversion:
    %(prog)s document.docx --mode convert --format markdown -o output.md
    %(prog)s results.json --mode convert --format docx -o output.docx

  Custom API Configuration:
    %(prog)s document.docx -o results.csv --api-key KEY --model gpt-4o
    %(prog)s document.docx -o results.csv --api-base https://api.example.com

OUTPUT FORMATS:
  .csv   - CSV with source filename (auto-creates .json for nested table preservation)
  .json  - JSON with complete structure and nested table markers
  .docx  - Word document with actual embedded nested tables (convert mode only)
  .md    - Markdown format (convert mode only)

ENVIRONMENT VARIABLES:
  Azure OpenAI (default):
    USE_AZURE_OPENAI=true              # Enable Azure OpenAI (default)
    AZURE_OPENAI_API_KEY=<key>         # API key
    AZURE_OPENAI_ENDPOINT=<url>        # Endpoint URL
    AZURE_OPENAI_API_VERSION=<ver>     # API version (default: 2024-08-01-preview)
    AZURE_OPENAI_DEPLOYMENT=<name>     # Deployment name (default: gpt-4o)

  Standard OpenAI:
    USE_AZURE_OPENAI=false             # Use standard OpenAI
    OPENAI_API_KEY=<key>               # API key
    OPENAI_API_BASE=<url>              # API base URL (optional)
    OPENAI_MODEL=<model>               # Model name (default: gpt-4o-2024-08-06)

NOTES:
  - Always uses smart batch processing to prevent output truncation
  - Automatically processes multiple rows together within token limits
  - CSV output automatically creates a JSON file for nested table preservation
  - Nested tables use markers: <<NESTED_TABLE_N_START>>...<<NESTED_TABLE_N_END>>
  - Use convert mode to transform JSON back to DOCX with actual nested tables
  - Only the main 5-column table is processed; other tables in the document are ignored
        """
    )

    # Positional argument
    parser.add_argument(
        "input",
        type=str,
        nargs='?',
        default=None,
        metavar="INPUT",
        help="Input file or folder path (DOCX for extraction, JSON for conversion)"
    )

    # Output options
    output_group = parser.add_argument_group('Output Options')
    output_group.add_argument(
        "-o", "--output",
        type=str,
        metavar="FILE",
        help="Output file path (.csv, .json, .docx, or .md)"
    )
    output_group.add_argument(
        "--stdout",
        action="store_true",
        help="Print JSON output to stdout (in addition to file output)"
    )

    # Operation mode
    mode_group = parser.add_argument_group('Operation Mode')
    mode_group.add_argument(
        "--mode",
        type=str,
        choices=['extract', 'convert'],
        default='extract',
        help="Operation mode: 'extract' (AI extraction) or 'convert' (format conversion). Default: extract"
    )
    mode_group.add_argument(
        "--format",
        type=str,
        choices=['markdown', 'docx'],
        metavar="FORMAT",
        help="Target format for convert mode: 'markdown' or 'docx'"
    )

    # Processing options
    process_group = parser.add_argument_group('Processing Options')
    process_group.add_argument(
        "--batch-size",
        type=int,
        default=10,
        metavar="N",
        help="Max rows per batch (default: 10). Actual batch size may be smaller due to token limits."
    )
    process_group.add_argument(
        "--concurrency",
        type=int,
        default=5,
        metavar="N",
        help="Max concurrent API calls (default: 5)"
    )
    process_group.add_argument(
        "--max-tokens",
        type=int,
        default=20000,
        metavar="N",
        help="Max input tokens per batch (default: 20000). Controls smart batching behavior."
    )

    # API configuration
    api_group = parser.add_argument_group('API Configuration')
    api_group.add_argument(
        "--api-key",
        type=str,
        metavar="KEY",
        help="API key (overrides environment variables)"
    )
    api_group.add_argument(
        "--api-base",
        type=str,
        metavar="URL",
        help="API endpoint URL (overrides environment variables)"
    )
    api_group.add_argument(
        "--model",
        type=str,
        metavar="NAME",
        help="Model or deployment name (overrides environment variables)"
    )

    return parser.parse_args()


def main():
    """Main entry point for the CLI tool.

    Parses command line arguments, validates input, and processes DOCX files to
    extract structured information. Supports single file processing and batch
    folder processing. Results are printed to standard output and can optionally
    be saved to JSON or CSV files.

    Returns:
        Exit code (0 for success, 1 for failure).
    """
    args = parse_args()

    # Validate input
    if not args.input:
        print("Error: INPUT is required", file=sys.stderr)
        print("Run 'python main.py --help' for usage information", file=sys.stderr)
        return 1

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Path does not exist - {args.input}", file=sys.stderr)
        return 1

    # ========== CONVERT MODE ==========
    if args.mode == 'convert':
        if not args.format:
            print("Error: --format is required when using --mode convert", file=sys.stderr)
            print("Available formats: markdown, docx", file=sys.stderr)
            return 1

        if not args.output:
            print("Error: --output (-o) is required when using --mode convert", file=sys.stderr)
            return 1

        # Convert DOCX to Markdown
        if args.format == 'markdown':
            if not input_path.is_file():
                print("Error: Markdown conversion only supports single file", file=sys.stderr)
                return 1

            if input_path.suffix.lower() not in ['.docx', '.doc']:
                print(f"Warning: File may not be in DOCX format - {args.input}", file=sys.stderr)

            try:
                # Create a temporary extractor just for conversion (no API needed)
                temp_extractor = DocxExtractor(
                    api_key="dummy",  # Not used for markdown conversion
                    use_azure=False
                )

                print(f"Converting tables to Markdown: {input_path}")
                markdown_content = temp_extractor.convert_tables_to_markdown(str(input_path))

                # Save to file
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                print(f"Markdown content saved to: {args.output}")
                return 0
            except Exception as e:
                print(f"Error: Failed to convert to Markdown - {e}", file=sys.stderr)
                import traceback
                traceback.print_exc()
                return 1

        # Convert JSON to DOCX
        elif args.format == 'docx':
            if input_path.suffix.lower() != '.json':
                print("Error: DOCX conversion requires JSON input file", file=sys.stderr)
                return 1

            try:
                print(f"Converting JSON to DOCX: {input_path}")
                json_to_docx(str(input_path), args.output)
                print("Conversion complete!")
                return 0
            except Exception as e:
                print(f"Error: Failed to convert JSON to DOCX - {e}", file=sys.stderr)
                import traceback
                traceback.print_exc()
                return 1

    # ========== EXTRACT MODE ==========
    # Check whether to use Azure OpenAI (default is true)
    use_azure = os.getenv("USE_AZURE_OPENAI", "true").lower() in ("true", "1", "yes")

    if use_azure:
        # Get Azure OpenAI configuration
        api_key = args.api_key or os.getenv("AZURE_OPENAI_API_KEY")
        azure_endpoint = args.api_base or os.getenv("AZURE_OPENAI_ENDPOINT")
        azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview")
        model = args.model or os.getenv("AZURE_OPENAI_DEPLOYMENT") or "gpt-4o"

        if not api_key:
            print("Error: Azure OpenAI API key not provided", file=sys.stderr)
            print("Please provide the API key in one of the following ways:", file=sys.stderr)
            print("  1. Use the --api-key parameter", file=sys.stderr)
            print("  2. Set the AZURE_OPENAI_API_KEY environment variable", file=sys.stderr)
            print("     Example: set AZURE_OPENAI_API_KEY=your-api-key-here", file=sys.stderr)
            return 1

        if not azure_endpoint:
            print("Error: Azure OpenAI endpoint not provided", file=sys.stderr)
            print("Please provide the endpoint in one of the following ways:", file=sys.stderr)
            print("  1. Use the --api-base parameter", file=sys.stderr)
            print("  2. Set the AZURE_OPENAI_ENDPOINT environment variable", file=sys.stderr)
            print("     Example: set AZURE_OPENAI_ENDPOINT=https://your-resource.openai.azure.com/", file=sys.stderr)
            return 1
    else:
        # Get standard OpenAI configuration
        api_key = args.api_key or os.getenv("OPENAI_API_KEY")
        api_base = args.api_base or os.getenv("OPENAI_API_BASE")
        model = args.model or os.getenv("OPENAI_MODEL") or "gpt-4o-2024-08-06"
        azure_endpoint = None
        azure_api_version = None

        if not api_key:
            print("Error: OpenAI API key not provided", file=sys.stderr)
            print("Please provide the API key in one of the following ways:", file=sys.stderr)
            print("  1. Use the --api-key parameter", file=sys.stderr)
            print("  2. Set the OPENAI_API_KEY environment variable", file=sys.stderr)
            print("     Example: export OPENAI_API_KEY='your-api-key-here'", file=sys.stderr)
            return 1

    # Create extractor instance
    if use_azure:
        extractor = DocxExtractor(
            api_key=api_key,
            model=model,
            use_azure=True,
            azure_endpoint=azure_endpoint,
            azure_api_version=azure_api_version
        )
    else:
        extractor = DocxExtractor(
            api_key=api_key,
            model=model,
            api_base=api_base,
            use_azure=False
        )

    try:
        # Determine whether it's a file or folder
        if input_path.is_file():
            # Process single file
            if input_path.suffix.lower() not in ['.docx', '.doc']:
                print(f"Warning: File may not be in DOCX format - {args.input}", file=sys.stderr)

            # Always use batch processing (smart batching)
            extraction = extractor.process_file_with_batches(
                str(input_path),
                batch_size=args.batch_size,
                max_concurrent=args.concurrency,
                max_tokens_per_batch=args.max_tokens,
                output_path=args.output
            )

            # Output JSON format if --stdout flag is specified
            if args.stdout:
                print("\n" + "=" * 80)
                print("JSON output:")
                print(json.dumps(extraction.model_dump(), ensure_ascii=False, indent=2))

        elif input_path.is_dir():
            # Process all DOCX files in the folder
            docx_files = list(input_path.glob("*.docx")) + list(input_path.glob("*.doc"))

            if not docx_files:
                print(f"Error: No DOCX files found in folder - {args.input}", file=sys.stderr)
                return 1

            print(f"Found {len(docx_files)} DOCX file(s)")
            print(f"Processing settings: {args.batch_size} max rows/batch, {args.concurrency} max concurrent calls")
            print("=" * 80)

            extractions = []
            for idx, docx_file in enumerate(docx_files, 1):
                print(f"\n[{idx}/{len(docx_files)}] Processing file: {docx_file.name}")
                print("-" * 80)

                try:
                    extraction = extractor.process_file_with_batches(
                        str(docx_file),
                        batch_size=args.batch_size,
                        max_concurrent=args.concurrency,
                        max_tokens_per_batch=args.max_tokens
                    )
                    extractions.append((docx_file.name, extraction))
                except Exception as e:
                    print(f"Warning: Error processing file {docx_file.name}: {e}", file=sys.stderr)
                    continue

            # Save results
            if args.output:
                output_path = Path(args.output)
                if output_path.suffix.lower() == '.csv':
                    # Export to CSV
                    DocxExtractor.export_to_csv(extractions, str(output_path))
                    print(f"\nAll results saved to CSV file: {output_path}")

                    # Also save JSON with same base name
                    json_path = output_path.with_suffix('.json')
                    all_data = {
                        "files": [
                            {
                                "filename": filename,
                                "records": extraction.model_dump()["records"]
                            }
                            for filename, extraction in extractions
                        ]
                    }
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump(all_data, f, ensure_ascii=False, indent=2)
                    print(f"All results also saved to JSON file: {json_path}")

                elif output_path.suffix.lower() == '.docx':
                    # Export to DOCX - create separate files for each source
                    for filename, extraction in extractions:
                        base_name = Path(filename).stem
                        output_file = output_path.parent / f"{output_path.stem}_{base_name}.docx"
                        DocxExtractor.export_to_docx(extraction, str(output_file), filename)
                        print(f"Exported {filename} to {output_file}")
                    print(f"\nAll results saved to DOCX files in: {output_path.parent}")
                elif output_path.suffix.lower() == '.json':
                    # Export to JSON
                    all_data = {
                        "files": [
                            {
                                "filename": filename,
                                "records": extraction.model_dump()["records"]
                            }
                            for filename, extraction in extractions
                        ]
                    }
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(all_data, f, ensure_ascii=False, indent=2)
                    print(f"\nAll results saved to JSON file: {output_path}")
                else:
                    print(f"Warning: Unsupported output format {output_path.suffix}, please use .csv, .json, or .docx", file=sys.stderr)

            # Output JSON format to standard output if --stdout flag is specified
            if args.stdout:
                all_data = {
                    "files": [
                        {
                            "filename": filename,
                            "records": extraction.model_dump()["records"]
                        }
                        for filename, extraction in extractions
                    ]
                }
                print("\n" + "=" * 80)
                print("JSON output:")
                print(json.dumps(all_data, ensure_ascii=False, indent=2))

        else:
            print(f"Error: Input path is neither a file nor a folder - {args.input}", file=sys.stderr)
            return 1

        return 0

    except FileNotFoundError as e:
        print(f"Error: File not found - {e}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"Error: An error occurred during processing - {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())